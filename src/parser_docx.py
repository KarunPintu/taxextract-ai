from io import BytesIO
from typing import Dict

try:
    from docx import Document
except Exception:  # pragma: no cover - handled at runtime in the app
    Document = None

from src.utils import clean_text


def extract_docx_text(file_bytes: bytes) -> Dict[str, object]:
    result = {"text": "", "paragraphs": 0, "error": "", "parser": "python-docx"}
    if Document is None:
        result["error"] = "python-docx is not available."
        return result

    try:
        document = Document(BytesIO(file_bytes))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
        result["paragraphs"] = len(parts)
        result["text"] = clean_text("\n".join(parts))
    except Exception as exc:
        result["error"] = f"DOCX parsing failed: {exc}"
    return result
